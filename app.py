import os
import base64
from flask import Flask, request, render_template, jsonify, send_from_directory, Response
from werkzeug.utils import secure_filename
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, PyMuPDFLoader, UnstructuredExcelLoader, CSVLoader, WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.vectorstores import Chroma
import pandas as pd
import json
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDF preview
import docx
from langchain.docstore.document import Document
import io
from PIL import Image
import time
import gc
from llama_cloud_services import LlamaParse
from llama_index.core import SimpleDirectoryReader
from langchain.prompts import PromptTemplate
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup
import sqlite3
from config import (
    GOOGLE_API_KEY,
    GEMINI_MODEL,
    EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    RETRIEVER_K,
    INDEX_PERSIST_DIRECTORY,
    COLLECTION_NAME
)
from utils import save_chat_history, load_chat_history, format_chat_message
import uuid
import shutil

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx', 'csv'}

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(INDEX_PERSIST_DIRECTORY, exist_ok=True)

# Database initialization
def init_db():
    try:
        print("Initializing database...")
        conn = sqlite3.connect('documents.db')
        c = conn.cursor()
        
        # Create agents table if it doesn't exist
        c.execute('''
            CREATE TABLE IF NOT EXISTS agents
            (agent_id TEXT PRIMARY KEY,
             name TEXT NOT NULL,
             chroma_collection TEXT NOT NULL,
             created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)
        ''')
        
        # Create documents table with agent_id foreign key if it doesn't exist
        c.execute('''
            CREATE TABLE IF NOT EXISTS documents
            (doc_id TEXT PRIMARY KEY,
             agent_id TEXT NOT NULL,
             filename TEXT NOT NULL,
             chunk_count INTEGER NOT NULL,
             chroma_ids TEXT NOT NULL,
             created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
             FOREIGN KEY (agent_id) REFERENCES agents(agent_id))
        ''')
        
        conn.commit()
        conn.close()
        print("Database initialized successfully")
    except Exception as e:
        print(f"Error initializing database: {str(e)}")
        raise

# Initialize database
init_db()

# Database helper functions
def save_agent_to_db(agent_id, name, chroma_collection):
    try:
        conn = sqlite3.connect('documents.db')
        c = conn.cursor()
        
        # Check if agent already exists
        c.execute('SELECT * FROM agents WHERE agent_id = ?', (agent_id,))
        existing_agent = c.fetchone()
        
        if existing_agent:
            print(f"Updating existing agent: {agent_id}")
            c.execute('''
                UPDATE agents 
                SET name = ?, chroma_collection = ?
                WHERE agent_id = ?
            ''', (name, chroma_collection, agent_id))
        else:
            print(f"Creating new agent: {agent_id}")
            c.execute('''
                INSERT INTO agents (agent_id, name, chroma_collection)
                VALUES (?, ?, ?)
            ''', (agent_id, name, chroma_collection))
            
        conn.commit()
        conn.close()
        print(f"Agent {agent_id} saved successfully")
    except Exception as e:
        print(f"Error saving agent to database: {str(e)}")
        raise

def get_agent_from_db(agent_id):
    conn = sqlite3.connect('documents.db')
    c = conn.cursor()
    c.execute('SELECT * FROM agents WHERE agent_id = ?', (agent_id,))
    agent = c.fetchone()
    conn.close()
    if agent:
        return {
            'agent_id': agent[0],
            'name': agent[1],
            'created_at': agent[2]
        }
    return None

def save_document_to_db(doc_id, agent_id, filename, chunk_count, chroma_ids):
    conn = sqlite3.connect('documents.db')
    c = conn.cursor()
    c.execute('''
        INSERT OR REPLACE INTO documents (doc_id, agent_id, filename, chunk_count, chroma_ids)
        VALUES (?, ?, ?, ?, ?)
    ''', (doc_id, agent_id, filename, chunk_count, json.dumps(chroma_ids)))
    conn.commit()
    conn.close()

def get_document_from_db(doc_id):
    conn = sqlite3.connect('documents.db')
    c = conn.cursor()
    c.execute('''
        SELECT d.*, a.name as agent_name 
        FROM documents d 
        JOIN agents a ON d.agent_id = a.agent_id 
        WHERE d.doc_id = ?
    ''', (doc_id,))
    doc = c.fetchone()
    conn.close()
    if doc:
        return {
            'doc_id': doc[0],
            'agent_id': doc[1],
            'filename': doc[2],
            'chunk_count': doc[3],
            'chroma_ids': json.loads(doc[4]),
            'created_at': doc[5],
            'agent_name': doc[6]
        }
    return None

def get_all_documents(agent_id=None):
    conn = sqlite3.connect('documents.db')
    c = conn.cursor()
    try:
        if agent_id:
            print(f"Fetching documents for agent: {agent_id}")
            c.execute('''
                SELECT d.*, a.name as agent_name 
                FROM documents d 
                JOIN agents a ON d.agent_id = a.agent_id 
                WHERE d.agent_id = ?
            ''', (agent_id,))
        else:
            print("No agent ID provided, fetching all documents")
            c.execute('''
                SELECT d.*, a.name as agent_name 
                FROM documents d 
                JOIN agents a ON d.agent_id = a.agent_id
            ''')
        docs = c.fetchall()
        print(f"Found {len(docs)} documents in database")
        return [{
            'doc_id': doc[0],
            'agent_id': doc[1],
            'filename': doc[2],
            'chunk_count': doc[3],
            'chroma_ids': json.loads(doc[4]),
            'created_at': doc[5],
            'agent_name': doc[6]
        } for doc in docs]
    except Exception as e:
        print(f"Error in get_all_documents: {str(e)}")
        raise
    finally:
        conn.close()

def delete_document_from_db(doc_id):
    conn = sqlite3.connect('documents.db')
    c = conn.cursor()
    c.execute('DELETE FROM documents WHERE doc_id = ?', (doc_id,))
    conn.commit()
    conn.close()

def delete_agent_from_db(agent_id):
    conn = sqlite3.connect('documents.db')
    c = conn.cursor()
    # First delete all documents associated with the agent
    c.execute('DELETE FROM documents WHERE agent_id = ?', (agent_id,))
    # Then delete the agent
    c.execute('DELETE FROM agents WHERE agent_id = ?', (agent_id,))
    conn.commit()
    conn.close()

# Global variables for document processing
conversation_chain = None
vector_store = None  # Initialize vector_store as None
current_agent_id = None  # Track current agent

# Initialize Gemini
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")
LLAMA_PARSE_KEY = os.getenv('LLAMA_PARSE_KEY')
if not LLAMA_PARSE_KEY:
    raise ValueError("LLAMA_PARSE_KEY environment variable is not set")

# Initialize embeddings
embeddings = GoogleGenerativeAIEmbeddings(
    model=EMBEDDING_MODEL,
    google_api_key=GOOGLE_API_KEY
)

# Initialize LLM
llm = ChatGoogleGenerativeAI(
    model=GEMINI_MODEL,
    google_api_key=GOOGLE_API_KEY,
    temperature=0.7,
    convert_system_message_to_human=True
)

# Initialize memory
memory = ConversationBufferMemory(
    memory_key="chat_history",
    return_messages=True,
    output_key="answer"
)

# Initialize Chroma vector store
try:
    vector_store = Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=embeddings,
        persist_directory=INDEX_PERSIST_DIRECTORY
    )
except Exception as e:
    print(f"Error loading existing Chroma store: {str(e)}")
    # Create a new store if loading fails
    dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy"})]
    vector_store = Chroma.from_documents(
        documents=dummy_docs,
        embedding=embeddings,
        collection_name=COLLECTION_NAME,
        persist_directory=INDEX_PERSIST_DIRECTORY
    )

# Initialize conversation chain with retriever
conversation_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=vector_store.as_retriever(search_kwargs={"k": RETRIEVER_K}),
    memory=memory,
    return_source_documents=True,
    verbose=True,
    combine_docs_chain_kwargs={"prompt": PromptTemplate(
        input_variables=["context", "question", "chat_history"],
        template="""You are a helpful AI assistant that answers questions based on the provided context. 
        Use the following pieces of context to answer the question at the end. 
        If you don't know the answer, just say that you don't know, don't try to make up an answer.

        Context: {context}

        Chat History: {chat_history}

        Question: {question}

        Answer:"""
    )}
)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_document(file_path, doc_id):
    """Process document and create vector store"""
    global conversation_chain, vector_store, current_agent_id
    
    if not current_agent_id:
        raise RuntimeError("No active agent. Please create an agent first.")

    # Load document based on file type
    if file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith('.docx'):
        loader = Docx2txtLoader(file_path)
    elif file_path.endswith('.xlsx'):
        loader = UnstructuredExcelLoader(file_path)
    elif file_path.endswith('.csv'):
        loader = CSVLoader(file_path)
    else:
        raise ValueError("Unsupported file type")

    # Load and split the document
    raw_documents = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len
    )

    try:
        parser = LlamaParse(result_type="markdown", api_key=LLAMA_PARSE_KEY)
        file_extractor = {".pdf": parser}
        raw_docs = SimpleDirectoryReader(input_files=[file_path], file_extractor=file_extractor).load_data()
    except Exception as e:
        raise RuntimeError(f"LlamaParse failed: {str(e)}")

    langchain_docs = [
        Document(
            page_content=getattr(doc, 'text', ''),
            metadata={"source": file_path, "doc_id": doc_id, "agent_id": current_agent_id}
        )
        for doc in raw_docs if hasattr(doc, 'text') and doc.text
    ]
    splits = text_splitter.split_documents(langchain_docs)

    # Add documents to Chroma store and get their IDs
    try:
        # Get existing documents for this agent
        existing_docs = get_all_documents(current_agent_id)
        
        # If this is the first document, create a new vector store
        if not existing_docs:
            print("First document for this agent, creating new vector store...")
            vector_store = Chroma.from_documents(
                documents=splits,
                embedding=embeddings,
                collection_name=f'agent_{current_agent_id}',
                persist_directory=os.path.join('chroma_db', f'agent_{current_agent_id}')
            )
        else:
            print("Adding document to existing vector store...")
            # Add to existing vector store
            ids = vector_store.add_documents(splits)
            vector_store.persist()
        
        # Save document information to database
        save_document_to_db(
            doc_id=doc_id,
            agent_id=current_agent_id,
            filename=os.path.basename(file_path),
            chunk_count=len(splits),
            chroma_ids=ids if 'ids' in locals() else []
        )
        
        # Update conversation chain with new retriever
        print("Updating conversation chain with new documents...")
        retriever = vector_store.as_retriever(search_kwargs={"k": RETRIEVER_K})
        initialize_conversation_chain(retriever)
        print("Conversation chain updated successfully")
    except Exception as e:
        raise RuntimeError(f"Failed to add documents to Chroma store: {str(e)}")

    return splits

def get_website_title(url):
    """Extract title from website"""
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.string if soup.title else urlparse(url).netloc
        return title.strip()
    except:
        return urlparse(url).netloc

def process_url(url):
    """Process website URL and create document chunks"""
    global conversation_chain, vector_store, current_agent_id
    
    if not current_agent_id:
        raise RuntimeError("No active agent. Please create an agent first.")

    try:
        # Create a unique document ID
        doc_id = f'web_{int(time.time())}'

        # Get website title for display
        title = get_website_title(url)

        # Load website content with error handling
        try:
            loader = WebBaseLoader(url)
            raw_documents = loader.load()
        except Exception as e:
            raise RuntimeError(f"Failed to load website content: {str(e)}")

        if not raw_documents:
            raise RuntimeError("No content found at the specified URL")

        # Split the content
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len
        )

        splits = text_splitter.split_documents(raw_documents)

        if not splits:
            raise RuntimeError("No content chunks created from the website")

        # Add documents to Chroma store and get their IDs
        try:
            ids = vector_store.add_documents(splits)
            vector_store.persist()
            
            # Save document information to database
            save_document_to_db(
                doc_id=doc_id,
                agent_id=current_agent_id,
                filename=f"{title} ({urlparse(url).netloc})",
                chunk_count=len(splits),
                chroma_ids=ids
            )
        except Exception as e:
            raise RuntimeError(f"Failed to add documents to Chroma store: {str(e)}")

        # Update conversation chain with new retriever
        try:
            conversation_chain.retriever = vector_store.as_retriever(
                search_kwargs={"k": RETRIEVER_K}
            )
        except Exception as e:
            raise RuntimeError(f"Failed to update conversation chain: {str(e)}")

        return doc_id, splits

    except Exception as e:
        raise RuntimeError(f"Error processing URL: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):
        try:
            # Save the file
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Generate unique document ID
            doc_id = f'doc_{int(time.time())}'

            # Process document and get chunks
            chunks = process_document(file_path, doc_id)

            # Prepare chunk information
            chunk_info = []
            for i, chunk in enumerate(chunks):
                chunk_info.append({
                    'id': f'{doc_id}_chunk_{i + 1}',
                    'doc_id': doc_id,
                    'page': chunk.metadata.get('page', 'N/A'),
                    'content': chunk.page_content[:100] + '...'  # First 100 chars as preview
                })

            return jsonify({
                'filename': filename,
                'document_id': doc_id,
                'chunk_count': len(chunks),
                'chunks': chunk_info
            })

        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify({'error': 'File type not allowed'}), 400

@app.route('/chat', methods=['POST'])
def chat():
    global conversation_chain, memory
    if conversation_chain is None:
        return jsonify({'error': 'Please upload a document first'}), 400

    data = request.json
    question = data.get('question')
    if not question:
        return jsonify({'error': 'No question provided'}), 400

    try:
        # Get chat history from memory
        chat_history = memory.load_memory_variables({}).get("chat_history", [])

        # Invoke the chain with both question and chat history
        result = conversation_chain.invoke({
            "question": question,
            "chat_history": chat_history
        })

        response = result['answer']
        sources = []

        # Extract page numbers and document IDs from source documents
        for doc in result['source_documents']:
            if 'page' in doc.metadata and 'doc_id' in doc.metadata:
                doc_id = doc.metadata['doc_id']
                doc_info = get_document_from_db(doc_id)
                if doc_info:
                    sources.append({
                        'doc_id': doc_id,
                        'filename': doc_info['filename'],
                        'page': doc.metadata['page'],
                        'text': doc.page_content[:200] + '...' if len(doc.page_content) > 200 else doc.page_content
                    })

        return jsonify({
            'response': response,
            'sources': sources
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete_document', methods=['POST'])
def delete_document():
    global conversation_chain, vector_store
    try:
        data = request.json
        doc_id = data.get('document_id')

        if not doc_id:
            return jsonify({'error': 'No document ID provided'}), 400

        # Get document from database
        doc = get_document_from_db(doc_id)
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        # Verify document belongs to current agent
        if doc['agent_id'] != current_agent_id:
            return jsonify({'error': 'Document does not belong to current agent'}), 403

        # Delete from ChromaDB
        try:
            if doc['chroma_ids']:
                print(f"Deleting ChromaDB IDs for document {doc_id}: {doc['chroma_ids']}")
                vector_store.delete(ids=doc['chroma_ids'])
                vector_store.persist()
        except Exception as e:
            print(f"Error deleting from ChromaDB: {str(e)}")
            return jsonify({'error': f'Error deleting from ChromaDB: {str(e)}'}), 500

        # Delete from database
        delete_document_from_db(doc_id)

        # If no documents left for this agent, reset to dummy store
        if not get_all_documents(current_agent_id):
            dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy", "agent_id": current_agent_id})]
            vector_store = Chroma.from_documents(
                documents=dummy_docs,
                embedding=embeddings,
                collection_name=f'agent_{current_agent_id}',
                persist_directory=os.path.join('chroma_db', f'agent_{current_agent_id}')
            )
            vector_store.persist()
            conversation_chain.retriever = vector_store.as_retriever(
                search_kwargs={"k": RETRIEVER_K}
            )

        return jsonify({'message': 'Document deleted successfully'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete_all_documents', methods=['POST'])
def delete_all_documents():
    global conversation_chain, vector_store, memory
    try:
        if not current_agent_id:
            return jsonify({'error': 'No active agent'}), 400

        # Get all documents for current agent
        docs = get_all_documents(current_agent_id)
        
        # Delete all documents from ChromaDB
        for doc in docs:
            try:
                if doc['chroma_ids']:
                    print(f"Deleting ChromaDB IDs for document {doc['doc_id']}: {doc['chroma_ids']}")
                    vector_store.delete(ids=doc['chroma_ids'])
            except Exception as e:
                print(f"Error deleting document {doc['doc_id']} from ChromaDB: {str(e)}")
                continue

        # Persist ChromaDB changes
        vector_store.persist()

        # Delete all documents from database for current agent
        conn = sqlite3.connect('documents.db')
        c = conn.cursor()
        c.execute('DELETE FROM documents WHERE agent_id = ?', (current_agent_id,))
        conn.commit()
        conn.close()

        # Reset to dummy store for this agent
        dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy", "agent_id": current_agent_id})]
        vector_store = Chroma.from_documents(
            documents=dummy_docs,
            embedding=embeddings,
            collection_name=f'agent_{current_agent_id}',
            persist_directory=os.path.join('chroma_db', f'agent_{current_agent_id}')
        )
        vector_store.persist()
        conversation_chain.retriever = vector_store.as_retriever(
            search_kwargs={"k": RETRIEVER_K}
        )

        # Clear chat memory
        print("Clearing chat memory after deleting all documents...")
        memory.clear()

        return jsonify({'message': 'All documents deleted successfully'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/list_documents', methods=['GET'])
def list_documents():
    """Return list of all uploaded documents with their chunk information"""
    try:
        # Get documents for the current agent
        docs = get_all_documents(current_agent_id)
        print(f"Found {len(docs)} documents for agent {current_agent_id}")
        return jsonify([{
            'document_id': doc['doc_id'],
            'filename': doc['filename'],
            'chunk_count': doc['chunk_count']
        } for doc in docs])
    except Exception as e:
        print(f"Error listing documents: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/clear_chat', methods=['POST'])
def clear_chat():
    global memory
    try:
        # Clear the memory
        memory.clear()

        return jsonify({'message': 'Chat history cleared successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/process_url', methods=['POST'])
def handle_url():
    data = request.json
    url = data.get('url')

    if not url:
        return jsonify({'error': 'No URL provided'}), 400

    try:
        # Process the URL
        doc_id, chunks = process_url(url)

        # Get document info from database
        doc_info = get_document_from_db(doc_id)
        if not doc_info:
            raise RuntimeError("Document not found in database")

        # Prepare chunk information
        chunk_info = []
        for i, chunk in enumerate(chunks):
            chunk_info.append({
                'id': f'{doc_id}_chunk_{i + 1}',
                'doc_id': doc_id,
                'page': 'N/A',  # Web pages don't have page numbers
                'content': chunk.page_content[:100] + '...'  # First 100 chars as preview
            })

        return jsonify({
            'filename': doc_info['filename'],
            'document_id': doc_id,
            'chunk_count': len(chunks),
            'chunks': chunk_info
        })

    except Exception as e:
        print(f"Error processing URL: {str(e)}")
        return jsonify({'error': str(e)}), 500

def initialize_conversation_chain(retriever):
    """Initialize or reinitialize the conversation chain with a new retriever"""
    global conversation_chain
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        memory=memory,
        return_source_documents=True,
        verbose=True,
        combine_docs_chain_kwargs={"prompt": PromptTemplate(
            input_variables=["context", "question", "chat_history"],
            template="""You are a helpful AI assistant that answers questions based on the provided context. 
            Use the following pieces of context to answer the question at the end. 
            If you don't know the answer, just say that you don't know, don't try to make up an answer.

            Context: {context}

            Chat History: {chat_history}

            Question: {question}

            Answer:"""
        )}
    )

@app.route('/create_agent', methods=['POST'])
def create_agent():
    global current_agent_id, conversation_chain, vector_store
    print("\n=== Starting Agent Creation Process ===")
    try:
        data = request.json
        agent_name = data.get('name')
        
        if not agent_name:
            print("Error: No agent name provided")
            return jsonify({'error': 'Agent name is required'}), 400
            
        # Generate unique agent ID
        agent_id = f'agent_{int(time.time())}'
        
        print(f"\n1. Creating agent with ID: {agent_id} and name: {agent_name}")
        
        # Save agent to database
        try:
            print("\n2. Database Operations:")
            print("   - Connecting to database...")
            conn = sqlite3.connect('documents.db')
            c = conn.cursor()
            print("   - Inserting agent into database...")
            save_agent_to_db(agent_id, agent_name, f"agent_{agent_id}")
            print("   ✓ Database operations completed successfully")
        except Exception as db_error:
            print(f"\n❌ Database error: {str(db_error)}")
            return jsonify({'error': f'Database error: {str(db_error)}'}), 500
            
        current_agent_id = agent_id
        print(f"\n3. Current agent ID set to: {current_agent_id}")

        # Create new vector store for this agent
        try:
            print("\n4. Vector Store Operations:")
            print(f"   - Using collection name: {f'agent_{agent_id}'}")
            print(f"   - Using persist directory: {os.path.join('chroma_db', f'agent_{agent_id}')}")
            
            # Create dummy document with agent_id
            dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy", "agent_id": agent_id})]
            print("   - Created dummy document")
            
            # Initialize vector store
            print("   - Initializing ChromaDB vector store...")
            vector_store = Chroma.from_documents(
                documents=dummy_docs,
                embedding=embeddings,
                collection_name=f'agent_{agent_id}',
                persist_directory=os.path.join('chroma_db', f'agent_{agent_id}')
            )
            print("   - Vector store initialized")
            
            # Persist the store
            print("   - Persisting vector store...")
            vector_store.persist()
            print("   ✓ Vector store operations completed successfully")
        except Exception as vector_error:
            print(f"\n❌ Vector store error: {str(vector_error)}")
            return jsonify({'error': f'Vector store error: {str(vector_error)}'}), 500

        # Initialize conversation chain
        try:
            print("\n5. Conversation Chain Operations:")
            print("   - Creating retriever...")
            retriever = vector_store.as_retriever(search_kwargs={"k": RETRIEVER_K})
            print("   - Initializing conversation chain...")
            initialize_conversation_chain(retriever)
            print("   ✓ Conversation chain operations completed successfully")
        except Exception as chain_error:
            print(f"\n❌ Conversation chain error: {str(chain_error)}")
            return jsonify({'error': f'Conversation chain error: {str(chain_error)}'}), 500

        # Clear chat memory for new agent
        memory.clear()
        
        # Prepare success response
        response_data = {
            'message': 'Agent created successfully',
            'agent_id': str(agent_id),
            'name': str(agent_name)
        }
        
        print("\n=== Agent Creation Process Completed Successfully ===\n")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"\n❌ Error creating agent: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/list_agents', methods=['GET'])
def list_agents():
    """Return list of all agents with their documents"""
    try:
        conn = sqlite3.connect('documents.db')
        c = conn.cursor()
        
        # Get all agents with their document counts
        c.execute('''
            SELECT a.agent_id, a.name, COUNT(d.doc_id) as document_count
            FROM agents a
            LEFT JOIN documents d ON a.agent_id = d.agent_id
            GROUP BY a.agent_id, a.name
            ORDER BY a.created_at DESC
        ''')
        
        agents = c.fetchall()
        result = []
        
        for agent in agents:
            agent_id, name, doc_count = agent
            # Get documents for this agent
            c.execute('''
                SELECT doc_id, filename, chunk_count
                FROM documents
                WHERE agent_id = ?
            ''', (agent_id,))
            documents = c.fetchall()
            
            result.append({
                'agent_id': agent_id,
                'name': name,
                'document_count': doc_count,
                'documents': [{
                    'doc_id': doc[0],
                    'filename': doc[1],
                    'chunk_count': doc[2]
                } for doc in documents]
            })
        
        conn.close()
        return jsonify(result)
    except Exception as e:
        print(f"Error listing agents: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/switch_agent', methods=['POST'])
def switch_agent():
    """Switch to a different agent"""
    global current_agent_id, conversation_chain, vector_store, memory
    try:
        data = request.json
        agent_id = data.get('agent_id')
        
        if not agent_id:
            return jsonify({'error': 'Agent ID is required'}), 400
            
        # Get agent details
        agent = get_agent_from_db(agent_id)
        if not agent:
            return jsonify({'error': 'Agent not found'}), 404
            
        # Update current agent
        current_agent_id = agent_id
        
        # Clear chat memory
        print("Clearing chat memory for new agent...")
        memory.clear()
        
        # Update vector store for this agent
        try:
            # Create dummy document with agent_id
            dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy", "agent_id": agent_id})]
            
            # Initialize vector store
            vector_store = Chroma.from_documents(
                documents=dummy_docs,
                embedding=embeddings,
                collection_name=f'agent_{agent_id}',
                persist_directory=os.path.join('chroma_db', f'agent_{agent_id}')
            )
            vector_store.persist()
            
            # Update conversation chain
            conversation_chain.retriever = vector_store.as_retriever(
                search_kwargs={"k": RETRIEVER_K}
            )
        except Exception as e:
            print(f"Error updating vector store: {str(e)}")
            return jsonify({'error': f'Error updating vector store: {str(e)}'}), 500
            
        return jsonify({
            'message': 'Agent switched successfully',
            'agent_id': agent_id,
            'name': agent['name']
        })
    except Exception as e:
        print(f"Error switching agent: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/delete_agent', methods=['POST'])
def delete_agent():
    """Delete an agent and all its associated data"""
    global current_agent_id, conversation_chain, vector_store, memory
    try:
        data = request.json
        agent_id = data.get('agent_id')
        
        if not agent_id:
            return jsonify({'error': 'Agent ID is required'}), 400
            
        # Get agent details
        agent = get_agent_from_db(agent_id)
        if not agent:
            return jsonify({'error': 'Agent not found'}), 404
            
        print(f"Deleting agent: {agent_id}")
        
        # If we're deleting the current agent, reset the global variables first
        if current_agent_id == agent_id:
            # Reset vector store
            if vector_store:
                try:
                    vector_store._client.close()
                except:
                    pass
            vector_store = None
            current_agent_id = None
            memory.clear()
            
            # Create a temporary dummy vector store and conversation chain
            dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy"})]
            temp_vector_store = Chroma.from_documents(
                documents=dummy_docs,
                embedding=embeddings,
                collection_name="temp_store",
                persist_directory="temp_chroma"
            )
            initialize_conversation_chain(temp_vector_store.as_retriever(search_kwargs={"k": RETRIEVER_K}))
            print("Reset global variables before deletion")
        
        # Delete agent's Chroma collection
        try:
            chroma_dir = os.path.join('chroma_db', f'agent_{agent_id}')
            if os.path.exists(chroma_dir):
                # Force close any open file handles
                gc.collect()
                time.sleep(1)
                shutil.rmtree(chroma_dir, ignore_errors=True)
                print(f"Deleted Chroma collection: {chroma_dir}")
        except Exception as e:
            print(f"Error deleting Chroma collection: {str(e)}")
            # Continue with database deletion even if Chroma deletion fails
            
        # Delete agent and its documents from database
        try:
            conn = sqlite3.connect('documents.db')
            c = conn.cursor()
            c.execute('DELETE FROM documents WHERE agent_id = ?', (agent_id,))
            c.execute('DELETE FROM agents WHERE agent_id = ?', (agent_id,))
            conn.commit()
            conn.close()
            print("Deleted agent and documents from database")
        except Exception as e:
            print(f"Error deleting from database: {str(e)}")
            return jsonify({'error': f'Error deleting from database: {str(e)}'}), 500
            
        return jsonify({'message': 'Agent deleted successfully'})
        
    except Exception as e:
        print(f"Error deleting agent: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5006))
    app.run(host='0.0.0.0', port=port) 