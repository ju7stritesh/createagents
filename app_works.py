import os
import base64
from flask import Flask, request, render_template, jsonify, send_from_directory, Response
from werkzeug.utils import secure_filename
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, PyMuPDFLoader, UnstructuredExcelLoader, CSVLoader, WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain_google_genai import ChatGoogleGenerativeAI
import pandas as pd
import json
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDF preview
import docx
from langchain.docstore.document import Document
import io
from PIL import Image
import time
import gc
from llama_cloud_services import LlamaParse
from llama_index.core import SimpleDirectoryReader
from langchain.vectorstores import Chroma, FAISS
from langchain.prompts import PromptTemplate
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup
from config import (
    GOOGLE_API_KEY,
    GEMINI_MODEL,
    EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    RETRIEVER_K
)
from utils import save_chat_history, load_chat_history, format_chat_message

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx', 'csv'}

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize Gemini
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")
LLAMA_PARSE_KEY = os.getenv('LLAMA_PARSE_KEY')
if not LLAMA_PARSE_KEY:
    raise ValueError("LLAMA_PARSE_KEY environment variable is not set")

# Global variables for document processing
conversation_chain = None
documents = {}  # Dictionary to store document information
vector_store = None  # Initialize vector_store as None

# Initialize embeddings
embeddings = GoogleGenerativeAIEmbeddings(
    model=EMBEDDING_MODEL,
    google_api_key=GOOGLE_API_KEY
)

# Initialize LLM
llm = ChatGoogleGenerativeAI(
    model=GEMINI_MODEL,
    google_api_key=GOOGLE_API_KEY,
    temperature=0.7,
    convert_system_message_to_human=True
)

# Initialize memory
memory = ConversationBufferMemory(
    memory_key="chat_history",
    return_messages=True,
    output_key="answer"
)

# Create a dummy vector store for initial retriever
dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy"})]
dummy_vector_store = FAISS.from_documents(dummy_docs, embeddings)

# Initialize conversation chain with dummy retriever
conversation_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=dummy_vector_store.as_retriever(search_kwargs={"k": RETRIEVER_K}),
    memory=memory,
    return_source_documents=True,
    verbose=True,
    combine_docs_chain_kwargs={"prompt": PromptTemplate(
        input_variables=["context", "question", "chat_history"],
        template="""You are a helpful AI assistant that answers questions based on the provided context. 
        Use the following pieces of context to answer the question at the end. 
        If you don't know the answer, just say that you don't know, don't try to make up an answer.

        Context: {context}

        Chat History: {chat_history}

        Question: {question}

        Answer:"""
    )}
)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_document(file_path, doc_id):
    """Process document and create vector store"""
    global conversation_chain, documents, vector_store

    # Load document based on file type
    if file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith('.docx'):
        loader = Docx2txtLoader(file_path)
    elif file_path.endswith('.xlsx'):
        loader = UnstructuredExcelLoader(file_path)
    elif file_path.endswith('.csv'):
        loader = CSVLoader(file_path)
    else:
        raise ValueError("Unsupported file type")

    # Load and split the document
    raw_documents = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len
    )

    try:
        parser = LlamaParse(result_type="markdown", api_key=LLAMA_PARSE_KEY)
        file_extractor = {".pdf": parser}
        raw_docs = SimpleDirectoryReader(input_files=[file_path], file_extractor=file_extractor).load_data()
    except Exception as e:
        raise RuntimeError(f"LlamaParse failed: {str(e)}")

    langchain_docs = [
        Document(
            page_content=getattr(doc, 'text', ''),
            metadata={"source": file_path, "doc_id": doc_id}
        )
        for doc in raw_docs if hasattr(doc, 'text') and doc.text
    ]
    splits = text_splitter.split_documents(langchain_docs)

    # Store document information
    documents[doc_id] = {
        'filename': os.path.basename(file_path),
        'chunks': splits,
        'chunk_count': len(splits)
    }

    # Create or update vector store
    if vector_store is None:
        vector_store = FAISS.from_documents(splits, embeddings)
    else:
        vector_store.add_documents(splits)

    # Update conversation chain with new retriever
    conversation_chain.retriever = vector_store.as_retriever(
        search_kwargs={"k": RETRIEVER_K}
    )

    return splits

def get_website_title(url):
    """Extract title from website"""
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.string if soup.title else urlparse(url).netloc
        return title.strip()
    except:
        return urlparse(url).netloc

def process_url(url):
    """Process website URL and create document chunks"""
    global conversation_chain, documents, vector_store

    try:
        # Create a unique document ID
        doc_id = f'web_{int(time.time())}'

        # Get website title for display
        title = get_website_title(url)

        # Load website content with error handling
        try:
            loader = WebBaseLoader(url)
            raw_documents = loader.load()
        except Exception as e:
            raise RuntimeError(f"Failed to load website content: {str(e)}")

        if not raw_documents:
            raise RuntimeError("No content found at the specified URL")

        # Split the content
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len
        )

        splits = text_splitter.split_documents(raw_documents)

        if not splits:
            raise RuntimeError("No content chunks created from the website")

        # Store document information
        documents[doc_id] = {
            'filename': f"{title} ({urlparse(url).netloc})",
            'chunks': splits,
            'chunk_count': len(splits),
            'url': url
        }

        # Create or update vector store
        try:
            if vector_store is None:
                vector_store = FAISS.from_documents(splits, embeddings)
            else:
                vector_store.add_documents(splits)
        except Exception as e:
            raise RuntimeError(f"Failed to create/update vector store: {str(e)}")

        # Update conversation chain with new retriever
        try:
            conversation_chain.retriever = vector_store.as_retriever(
                search_kwargs={"k": RETRIEVER_K}
            )
        except Exception as e:
            raise RuntimeError(f"Failed to update conversation chain: {str(e)}")

        return doc_id, splits

    except Exception as e:
        # Clean up if something went wrong
        if doc_id in documents:
            del documents[doc_id]
        raise RuntimeError(f"Error processing URL: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):
        try:
            # Save the file
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Generate unique document ID
            doc_id = f'doc_{int(time.time())}'

            # Process document and get chunks
            chunks = process_document(file_path, doc_id)

            # Prepare chunk information
            chunk_info = []
            for i, chunk in enumerate(chunks):
                chunk_info.append({
                    'id': f'{doc_id}_chunk_{i + 1}',
                    'doc_id': doc_id,
                    'page': chunk.metadata.get('page', 'N/A'),
                    'content': chunk.page_content[:100] + '...'  # First 100 chars as preview
                })

            return jsonify({
                'filename': filename,
                'document_id': doc_id,
                'chunk_count': len(chunks),
                'chunks': chunk_info
            })

        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify({'error': 'File type not allowed'}), 400

@app.route('/chat', methods=['POST'])
def chat():
    global conversation_chain, memory
    if conversation_chain is None:
        return jsonify({'error': 'Please upload a document first'}), 400

    data = request.json
    question = data.get('question')
    if not question:
        return jsonify({'error': 'No question provided'}), 400

    try:
        # Get chat history from memory
        chat_history = memory.load_memory_variables({}).get("chat_history", [])

        # Invoke the chain with both question and chat history
        result = conversation_chain.invoke({
            "question": question,
            "chat_history": chat_history
        })

        response = result['answer']
        sources = []

        # Extract page numbers and document IDs from source documents
        for doc in result['source_documents']:
            if 'page' in doc.metadata and 'doc_id' in doc.metadata:
                doc_id = doc.metadata['doc_id']
                doc_info = documents.get(doc_id, {})
                sources.append({
                    'doc_id': doc_id,
                    'filename': doc_info.get('filename', 'Unknown'),
                    'page': doc.metadata['page'],
                    'text': doc.page_content[:200] + '...' if len(doc.page_content) > 200 else doc.page_content
                })

        return jsonify({
            'response': response,
            'sources': sources
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete_document', methods=['POST'])
def delete_document():
    global conversation_chain, documents, vector_store, memory
    try:
        data = request.json
        doc_id = data.get('document_id')

        if doc_id and doc_id in documents:
            # Remove document from documents dictionary
            del documents[doc_id]

            # Recreate vector store with remaining documents
            if documents:
                all_chunks = []
                for doc_info in documents.values():
                    all_chunks.extend(doc_info['chunks'])
                vector_store = FAISS.from_documents(all_chunks, embeddings)
                conversation_chain.retriever = vector_store.as_retriever(
                    search_kwargs={"k": RETRIEVER_K}
                )
            else:
                # If no documents left, reset to dummy store
                dummy_docs = [Document(page_content="Dummy document", metadata={"source": "dummy"})]
                dummy_vector_store = FAISS.from_documents(dummy_docs, embeddings)
                conversation_chain.retriever = dummy_vector_store.as_retriever(
                    search_kwargs={"k": RETRIEVER_K}
                )
                vector_store = None

            return jsonify({'message': 'Document deleted successfully'})
        else:
            return jsonify({'error': 'Document not found'}), 404

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/list_documents', methods=['GET'])
def list_documents():
    """Return list of all uploaded documents with their chunk information"""
    doc_list = []
    for doc_id, doc_info in documents.items():
        doc_list.append({
            'document_id': doc_id,
            'filename': doc_info['filename'],
            'chunk_count': doc_info['chunk_count'],
            'chunks': [
                {
                    'id': f'{doc_id}_chunk_{i + 1}',
                    'page': chunk.metadata.get('page', 'N/A')
                }
                for i, chunk in enumerate(doc_info['chunks'])
            ]
        })
    return jsonify(doc_list)

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/clear_chat', methods=['POST'])
def clear_chat():
    global memory
    try:
        # Clear the memory
        memory.clear()

        return jsonify({'message': 'Chat history cleared successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/process_url', methods=['POST'])
def handle_url():
    data = request.json
    url = data.get('url')

    if not url:
        return jsonify({'error': 'No URL provided'}), 400

    try:
        # Process the URL
        doc_id, chunks = process_url(url)

        # Prepare chunk information
        chunk_info = []
        for i, chunk in enumerate(chunks):
            chunk_info.append({
                'id': f'{doc_id}_chunk_{i + 1}',
                'doc_id': doc_id,
                'page': 'N/A',  # Web pages don't have page numbers
                'content': chunk.page_content[:100] + '...'  # First 100 chars as preview
            })

        return jsonify({
            'filename': documents[doc_id]['filename'],
            'document_id': doc_id,
            'chunk_count': len(chunks),
            'chunks': chunk_info
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5006))
    app.run(host='0.0.0.0', port=port) 